"""Extract structure from the existing paper."""
from docx import Document
from docx.shared import Pt
import json

PATH = r"C:/Users/evija/Downloads/Vibe Auto Research - paper.docx"
doc = Document(PATH)

# Inventory: paragraphs, headings, tables, styles
out_lines = []
out_lines.append("# Paper Extract\n")
out_lines.append(f"Total paragraphs: {len(doc.paragraphs)}\n")
out_lines.append(f"Total tables: {len(doc.tables)}\n")
out_lines.append(f"Total sections: {len(doc.sections)}\n\n")

# Style inventory
styles = [s.name for s in doc.styles]
out_lines.append(f"## Styles (total {len(styles)})\n")
for s in styles[:60]:
    out_lines.append(f"- {s}")
out_lines.append("")

# Walk paragraphs with style info
out_lines.append("\n## Paragraphs (with style annotations)\n")
heading_index = []
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    style = p.style.name if p.style else "?"
    if style.startswith("Heading") or style == "Title":
        heading_index.append((i, style, txt))
    if txt:
        out_lines.append(f"[{i:04d}][{style}] {txt}")
    else:
        out_lines.append(f"[{i:04d}][{style}] <blank>")

out_lines.append("\n\n## Heading index\n")
for idx, style, txt in heading_index:
    out_lines.append(f"  [{idx:04d}][{style}] {txt}")

# Tables
out_lines.append("\n\n## Tables\n")
for ti, t in enumerate(doc.tables):
    out_lines.append(f"\n### Table {ti}: rows={len(t.rows)} cols={len(t.columns)} style={t.style.name if t.style else '?'}")
    for ri, row in enumerate(t.rows[:8]):
        cells = [c.text.strip()[:80] for c in row.cells]
        out_lines.append(f"  R{ri}: {cells}")

with open(r"C:/Users/evija/Downloads/_paper_extract.md", "w", encoding="utf-8") as f:
    f.write("\n".join(out_lines))

print("paragraphs:", len(doc.paragraphs))
print("tables:", len(doc.tables))
print("headings:", len(heading_index))
print("style sample:", styles[:30])
